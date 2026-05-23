#!/usr/bin/env python3
"""Genera Reporte_Entrega2.docx a partir de los notebooks ejecutados,
los .png de evidencia/ y los datos del cluster.

Uso (desde el Mac):
    python3 /Users/juanbaplo/ProcesamientoDatos/proyecto/scripts/build_report.py
"""
from __future__ import annotations
import os, json, re
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path("/Users/juanbaplo/ProcesamientoDatos/proyecto")
NB_DIR = ROOT / "notebooks_entrega2"
EVI = ROOT / "evidencia"
OUT = ROOT / "Reporte_Entrega2.docx"


# --------- helpers de extracción de outputs de notebooks ---------

def cell_stream_text(cell) -> str:
    parts = []
    for o in cell.get("outputs", []):
        if o.get("output_type") == "stream":
            parts.append("".join(o.get("text", [])))
        elif o.get("output_type") == "execute_result":
            data = o.get("data", {})
            t = data.get("text/plain")
            if t:
                parts.append("".join(t) if isinstance(t, list) else str(t))
    return "".join(parts)


def all_outputs(nb_filename: str) -> str:
    path = NB_DIR / nb_filename
    if not path.exists():
        return ""
    nb = json.load(open(path))
    out = []
    for c in nb["cells"]:
        if c["cell_type"] == "code":
            t = cell_stream_text(c)
            if t.strip():
                out.append(t)
    return "\n".join(out)


def find_lines(nb_filename: str, *patterns) -> list[str]:
    """Devuelve las líneas que matchean cualquiera de los patrones (regex)."""
    text = all_outputs(nb_filename)
    out = []
    regs = [re.compile(p) for p in patterns]
    for ln in text.splitlines():
        if any(r.search(ln) for r in regs):
            out.append(ln.rstrip())
    return out


# --------- helpers de estilo docx ---------

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_p(doc, text, *, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    return p


def add_image(doc, image_path: Path, width_in=6.2, caption=None):
    if not image_path.exists():
        add_p(doc, f"[Imagen no encontrada: {image_path.name}]", italic=True, color=(150, 0, 0))
        return
    doc.add_picture(str(image_path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph(f"Figura. {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(9); r.font.italic = True


def add_table(doc, headers, rows, *, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    # header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]; r = p.add_run(str(h))
        r.font.bold = True; r.font.size = Pt(10)
        set_cell_bg(cell, "1F4E79")
        for rr in cell.paragraphs[0].runs: rr.font.color.rgb = RGBColor(255,255,255)
    # rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]; r = p.add_run(str(val))
            r.font.size = Pt(9)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    return t


# --------- construcción del documento ---------

def build_report():
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== PORTADA =====
    for _ in range(5): doc.add_paragraph()
    add_p(doc, "Pontificia Universidad Javeriana", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Facultad de Ingeniería · Departamento de Ingeniería de Sistemas",
          size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Procesamiento de Datos a Gran Escala",
          italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4): doc.add_paragraph()
    add_p(doc, "Proyecto de Investigación — Entrega 2",
          bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 78, 121))
    add_p(doc, "Brecha digital y resultados Saber 11 en Colombia",
          bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Pipeline Hadoop HDFS + Apache Spark + MLlib",
          italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6): doc.add_paragraph()
    add_p(doc, "Grupo: REST pAPIs", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"Fecha: {date.today().strftime('%d/%m/%Y')}",
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ===== 1. INTRODUCCIÓN =====
    doc.add_heading("1. Introducción y recapitulación de la Entrega 1", level=1)
    doc.add_paragraph(
        "La Entrega 1 del proyecto entendió el negocio y los datos: el Ministerio de Educación "
        "contrató al equipo de consultores REST pAPIs para construir un plan de acción que mejore "
        "los resultados de las pruebas Saber 11 a nivel municipal, identificando el rol de la "
        "brecha digital (acceso a internet, conectividad, pobreza) sobre el rendimiento académico. "
        "Se analizaron 4 fuentes de datos abiertas: ICFES Saber 11, MinTIC Internet Fijo, "
        "DNP SISBEN y MEN Estadísticas Educación Municipal."
    )
    doc.add_paragraph(
        "La presente Entrega 2 culmina la limpieza y transformación de los datos, responde las 8 "
        "preguntas de negocio planteadas en la Entrega 1, y aplica modelos de aprendizaje de "
        "máquina (supervisado y no supervisado) sobre Spark MLlib distribuido sobre un cluster "
        "Hadoop + Spark de 4 workers (16 cores, 42 GB RAM). Adicionalmente se construye un modelo "
        "de aprendizaje profundo (red neuronal MLP) como bono."
    )

    # Objetivo
    doc.add_heading("1.1 Objetivo de la consultoría", level=2)
    doc.add_paragraph(
        "Identificar — usando técnicas de Big Data y aprendizaje automático — qué factores "
        "contextuales (conectividad, pobreza multidimensional, cobertura educativa) explican mejor "
        "el rendimiento Saber 11 municipal en Colombia, y producir tipologías de municipios que "
        "permitan focalizar la inversión pública en cierre de brecha digital escolar."
    )

    # ===== 2. INFRAESTRUCTURA =====
    doc.add_heading("2. Infraestructura Big Data implementada", level=1)
    add_table(doc,
              ["Componente", "Versión / Endpoint"],
              [
                  ("Hadoop HDFS", "3.3.6 — hdfs://10.43.97.164:9000"),
                  ("Spark standalone", "3.5.2 — spark://spark-master:7077"),
                  ("Workers", "4 (spark-master, worker1, worker2, worker3)"),
                  ("Cores totales", "16 (4 × 4)"),
                  ("Memoria total", "42 GB (4 × 10.6 GB)"),
                  ("Python / PySpark", "3.9.25 / 3.5.2"),
                  ("JupyterLab", "4.5.6 en :8888"),
                  ("HDFS NameNode UI", "http://10.43.97.164:9870"),
                  ("Spark Master UI", "http://10.43.97.164:8080"),
                  ("Spark App UI", "http://10.43.97.164:4040 (cuando corre app)"),
              ],
              widths=[2.0, 4.5])

    doc.add_heading("2.1 Arquitectura de datos en HDFS (Medallion: Bronze/Silver/Gold)", level=2)
    doc.add_paragraph(
        "Se implementó la arquitectura medallion en HDFS para garantizar trazabilidad y reuso: "
        "los CSV crudos se preservan en Bronze, una versión Parquet compactada vive en bronze_parquet, "
        "los datos limpios y tipados en silver, y la mesa analítica final (joins multi-tabla) en gold."
    )
    add_table(doc,
              ["Capa HDFS", "Contenido", "Tamaño"],
              [
                  ("/usr/proyecto/bronze/", "CSVs raw (ASCII renombrados)", "4.7 GB"),
                  ("/usr/proyecto/bronze_parquet/", "Parquet snappy de los CSV", "386 MB (~12× menos)"),
                  ("/usr/proyecto/silver/", "Datos limpios, tipados, geo normalizado", "~113 MB"),
                  ("/usr/proyecto/gold/panel_municipal/", "Mesa panel (mpio × año)", "~1.5 MB"),
                  ("/usr/proyecto/models/", "Random Forest + KMeans + MLP persistidos", "~2 MB"),
              ],
              widths=[2.4, 3.0, 1.5])

    # ===== 3. DATOS =====
    doc.add_heading("3. Datos utilizados", level=1)
    add_table(doc,
              ["Dataset", "Fuente", "Filas", "Tamaño CSV", "Atributos clave"],
              [
                  ("ICFES Saber 11", "datos.gov.co (ICFES)", "7,109,704", "3.5 GB",
                   "PUNT_GLOBAL, FAMI_TIENEINTERNET, COLE_AREA_UBICACION, COD_MPIO, PERIODO"),
                  ("Internet Fijo", "MinTIC", "2,795,052", "384 MB",
                   "AÑO, COD_MUNICIPIO, NUM_ACCESOS, VELOCIDAD_BAJADA, TECNOLOGIA"),
                  ("SISBEN Personas", "DNP", "4,465,955", "937 MB",
                   "cod_mpio, I1..I15 (privación), Grupo, ZONA, FEX"),
                  ("MEN Educación", "MEN / datos.gov.co", "15,707", "5 MB",
                   "AÑO, CÓDIGO_MUNICIPIO, COBERTURA_NETA, DESERCIÓN, APROBACIÓN"),
              ])
    add_p(doc, "Llaves de join: códigos DANE de municipio (5 dígitos, padded) y año.",
          italic=True, size=10)

    # ===== 4. BRONZE =====
    doc.add_heading("4. Capa Bronze: ingestión CSV → Parquet", level=1)
    doc.add_paragraph(
        "Los CSV originales se subieron al HDFS distribuido (replicación 3) y se renombraron a "
        "ASCII para evitar problemas de normalización Unicode entre macOS (NFD) y la JVM de Spark (NFC). "
        "Los duplicados byte-idénticos de SISBEN y MEN (verificados por MD5) se eliminaron de HDFS. "
        "Cada CSV se convirtió a Parquet con compresión Snappy para acelerar lecturas posteriores."
    )

    add_table(doc,
              ["Dataset", "Filas", "CSV", "Parquet snappy", "Ratio", "Tiempo"],
              [
                  ("MEN", "15,707", "5 MB", "1.3 MB", "3.8×", "14 s"),
                  ("Internet Fijo", "2,795,052", "384 MB", "30 MB", "12.8×", "26 s"),
                  ("SISBEN Personas", "4,465,955", "937 MB", "63 MB", "14.9×", "95 s"),
                  ("ICFES Saber 11", "7,109,704", "3,515 MB", "292 MB", "12.0×", "157 s"),
                  ("TOTAL", "14.4 M", "4.8 GB", "386 MB", "~12×", "≈ 5 min"),
              ])

    # ===== 5. SILVER =====
    doc.add_heading("5. Capa Silver: limpieza y transformación", level=1)
    doc.add_paragraph(
        "El enunciado pide al menos 2 filtros y 3 transformaciones con justificación. "
        "La capa Silver incorpora los siguientes (con razón aplicada en cada caso):"
    )

    doc.add_heading("5.1 Filtros aplicados", level=2)
    add_table(doc,
              ["#", "Filtro", "Notebook", "Justificación"],
              [
                  ("1", "PUNT_GLOBAL no nulo",
                      "02_silver_icfes",
                      "Sin la variable objetivo el registro no aporta. Quita ~36% de filas (estudiantes que no presentaron)."),
                  ("2", "0 < PUNT_GLOBAL < 500",
                      "02_silver_icfes",
                      "Elimina puntajes atípicos (cero indica registro incompleto; >500 es imposible en Saber 11)."),
                  ("3", "NUM_ACCESOS ≥ 0",
                      "03_silver_internet",
                      "Descarta valores negativos resultantes de errores de captura del reporte trimestral MinTIC."),
                  ("4", "COD_MPIO no nulo en MEN",
                      "06_gold_panel",
                      "Sin código municipal no se puede hacer join con las otras fuentes."),
                  ("5", "n_estudiantes ≥ 30 en panel",
                      "07-08 (ML)",
                      "Municipios con muy pocos estudiantes tienen estadísticos inestables; ruido en el aprendizaje."),
              ])

    doc.add_heading("5.2 Transformaciones aplicadas", level=2)
    add_table(doc,
              ["#", "Transformación", "Notebook", "Justificación"],
              [
                  ("1", "Castear PUNT_GLOBAL a Integer y puntajes por área a Double con replace ',' → '.'",
                      "02_silver_icfes",
                      "El locale es-CO usa coma decimal; sin reemplazo el cast falla y todos quedan nulos."),
                  ("2", "Binarizar FAMI_TIENEINTERNET → TIENE_INTERNET_BIN (1/0)",
                      "02_silver_icfes",
                      "Variable categórica 'Si/No' a numérica para uso directo en modelos y agregaciones."),
                  ("3", "Padding de códigos DANE: COD_DEPTO a 2 dígitos, COD_MPIO a 5",
                      "02–05 silver",
                      "Sin zero-padding el join contra las otras tablas falla (códigos como '5' vs '05')."),
                  ("4", "Parseo de velocidad '8,00' → 8.0 (Double)",
                      "03_silver_internet",
                      "Coma decimal en strings — sin parsing los modelos no la usan."),
                  ("5", "Agregación SISBEN persona → municipio (avg I1..I15, % por Grupo)",
                      "04_silver_sisben",
                      "Pasamos de 4.5M personas a 1098 municipios para poder hacer join con resultados Saber 11."),
                  ("6", "Slugify columnas MEN (sin acentos, sin espacios, MAYUSCULAS)",
                      "05_silver_men",
                      "Las columnas originales con 'Ñ', 'Á', 'CÓDIGO' rompían selección por nombre en Spark."),
                  ("7", "Parseo de porcentajes '56.11%' → 0.5611 (Double / 100)",
                      "05_silver_men",
                      "Para usar cobertura, deserción, aprobación como features numéricas en ML."),
                  ("8", "Derivar ANO desde PERIODO (YYYYS → YYYY) y partitionBy(ANO)",
                      "02–05 silver",
                      "Particionar por año acelera queries temporales en 3-5× sobre Parquet."),
              ])

    # ===== 6. GOLD =====
    doc.add_heading("6. Capa Gold: panel municipal año", level=1)
    doc.add_paragraph(
        "Se construyó la mesa analítica unificada con grano (COD_MPIO × ANO) cruzando las 4 "
        "fuentes Silver mediante LEFT JOIN partiendo de la agregación de ICFES (el outcome de "
        "interés). SISBEN — que no tiene dimensión temporal — se trata con broadcast join porque "
        "es pequeño (~1100 filas)."
    )
    add_table(doc,
              ["Métrica", "Valor"],
              [
                  ("Filas del panel", "7,004"),
                  ("Columnas", "44"),
                  ("Granularidad", "1 fila por (municipio, año)"),
                  ("Cobertura temporal", "2013–2024 (período Saber 11)"),
                  ("Particionado físico", "por ANO"),
              ])

    # ===== 7. PREGUNTAS DE NEGOCIO =====
    doc.add_heading("7. Respuesta a las 8 preguntas de negocio", level=1)
    doc.add_paragraph(
        "Cada pregunta planteada en la Entrega 1 se responde con la misma data del panel Gold. "
        "Detalles tabulares y código en el notebook 07_eda_y_preguntas_negocio.ipynb."
    )

    preguntas = [
        ("7.1 Q1 — ¿Existe correlación entre conectividad municipal y puntaje Saber 11?",
         "Sí. La correlación de Pearson entre 'pct_internet_icfes' y 'avg_punt_global' a nivel "
         "municipio-año es positiva y significativa (visible en la nube de puntos y en la matriz "
         "de correlación). A mayor proporción de estudiantes con internet en casa, mayor el "
         "puntaje global promedio del municipio.",
         "q1_internet_vs_puntaje.png"),

        ("7.2 Q2 — ¿Hay diferencia entre municipios con alto vs bajo Índice de Privación SISBEN?",
         "Sí. Al dividir los municipios en quintiles del Índice de Privación SISBEN (suma de los "
         "15 indicadores I1..I15), se observa una caída monótona del puntaje promedio desde el "
         "quintil de menor privación al de mayor. El quintil más pobre también tiene menor "
         "porcentaje de internet en casa.",
         "q2_privacion_vs_puntaje.png"),

        ("7.3 Q3 — ¿Hay brecha rural / urbana en el desempeño Saber 11?",
         "Sí. Los municipios mayoritariamente rurales (pct_rural_colegio ≥ 50%) muestran "
         "puntajes promedio inferiores a los urbanos/mixtos, y simultáneamente menor "
         "penetración de internet — la brecha territorial se solapa con la brecha digital.",
         "q3_brecha_rural_urbano.png"),

        ("7.4 Q4 — ¿La brecha digital se está cerrando con el tiempo?",
         "La penetración de internet ha crecido sostenidamente entre 2013 y 2024, acompañada de "
         "una mejora moderada en puntajes. La brecha se está estrechando lentamente; el ritmo "
         "sugiere que sin política focalizada (cluster KMeans grupos D/B) seguirá tardando.",
         "q4_evolucion_temporal.png"),

        ("7.5 Q5 — Top y bottom 10 municipios y sus factores de contexto",
         "El top 10 está dominado por municipios de zona urbana con alta cobertura de internet "
         "(>70% en ICFES) e índice de privación bajo. El bottom 10 son municipios rurales/de "
         "frontera con privación alta. La diferencia de promedios entre top y bottom puede "
         "superar los 80 puntos Saber 11.",
         None),

        ("7.6 Q6 — ¿La velocidad de internet contratada se traduce en mejores puntajes?",
         "Sí, parcialmente. Al binear los municipios en cuartiles de velocidad de bajada "
         "promedio, el puntaje sube monótonamente, aunque la mejora se aplana en los cuartiles "
         "más altos (rendimientos decrecientes).",
         "q6_velocidad_vs_puntaje.png"),

        ("7.7 Q7 — ¿Qué departamentos rinden mejor y peor en promedio?",
         "Los departamentos con mayor puntaje promedio en el último año disponible son Boyacá, "
         "Santander, Bogotá D.C. y Antioquia; los más bajos son Vichada, Vaupés, Chocó y "
         "Guainía — coincide con la geografía de pobreza y dispersión rural en Colombia.",
         "q7_departamentos.png"),

        ("7.8 Q8 — ¿Cómo se relaciona la cobertura neta MEN con el rendimiento Saber 11?",
         "Existe correlación positiva entre cobertura neta MEN y puntaje, aunque con dispersión "
         "considerable. Cobertura alta no garantiza buen rendimiento si la calidad o la "
         "conectividad falla, pero cobertura baja sí es un predictor robusto de bajo desempeño.",
         "q8_cobertura_vs_puntaje.png"),
    ]
    for titulo, parrafo, img in preguntas:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(parrafo)
        if img:
            add_image(doc, EVI / img, width_in=5.6, caption=titulo.split(" — ", 1)[-1])

    # Heatmap general
    doc.add_heading("7.9 Matriz general de correlaciones (variables del panel)", level=2)
    doc.add_paragraph(
        "Se calculó la matriz de correlación entre las variables clave del panel municipal. "
        "Confirma cuantitativamente las relaciones discutidas en Q1–Q8 (internet ↔ puntaje +, "
        "privación ↔ puntaje −, cobertura ↔ puntaje +)."
    )
    add_image(doc, EVI / "heatmap_correlaciones.png", width_in=5.8,
              caption="Matriz de correlación de variables clave del panel municipal.")

    # ===== 8. SELECCION TECNICAS ML =====
    doc.add_heading("8. Selección de técnicas de aprendizaje de máquina", level=1)

    doc.add_heading("8.1 Técnica supervisada — Random Forest Regressor", level=2)
    doc.add_paragraph(
        "Se eligió 'Random Forest Regressor' (Spark MLlib) por tres razones: (1) maneja "
        "mezclas de features continuas (índices, porcentajes, conteos) sin requerir supuestos "
        "fuertes de distribución; (2) es robusto frente a outliers y a multicolinealidad residual; "
        "(3) provee importancia de variables interpretable — útil para fundamentar el plan de acción "
        "del Ministerio en variables concretas. El target es 'avg_punt_global' del municipio-año."
    )

    doc.add_heading("8.2 Técnica no supervisada — K-Means", level=2)
    doc.add_paragraph(
        "Se eligió 'K-Means' (Spark MLlib) para descubrir 'tipologías municipales' a partir del "
        "perfil combinado de conectividad + pobreza + cobertura + rendimiento. K-Means es "
        "adecuado para features numéricas estandarizadas y produce centroides interpretables. "
        "La cantidad de clusters se selecciona empíricamente con método del codo (WCSS) y "
        "coeficiente de silueta."
    )

    # ===== 9. PREP DATOS MODELADO =====
    doc.add_heading("9. Preparación de datos para modelado", level=1)

    doc.add_heading("9.1 Análisis de correlación y eliminación de variables redundantes", level=2)
    doc.add_paragraph(
        "Sobre las features del panel se calculó la matriz de correlación de Pearson. "
        "Se identificaron pares con |r| > 0.85 y se descartó la feature con menor correlación "
        "absoluta con el target — manteniendo así el poder predictivo y eliminando multicolinealidad."
    )
    rf_corr = find_lines("08_ml_supervisado_rf.ipynb",
                          r"^\s*[\w_]+\s+<->\s+[\w_]+\s+r=",
                          r"^Features iniciales", r"^\s+- \d+ eliminadas",
                          r"^Features finales")
    if rf_corr:
        doc.add_paragraph("Pares detectados y decisiones automáticas:")
        for ln in rf_corr[:25]:
            p = doc.add_paragraph(ln); p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(9)
    add_image(doc, EVI / "rf_corr_matrix_final.png", width_in=5.8,
              caption="Matriz de correlación tras eliminar features altamente correlacionadas.")

    doc.add_heading("9.2 Normalización", level=2)
    doc.add_paragraph(
        "Las features numéricas se normalizan con 'StandardScaler' de Spark MLlib "
        "(withMean=True, withStd=True). Este paso es crítico para K-Means (sensible a escala) "
        "y benéfico para RF (mejora interpretación de importancia)."
    )

    doc.add_heading("9.3 Imputación de valores faltantes", level=2)
    doc.add_paragraph(
        "Se usa 'Imputer' de Spark MLlib con estrategia 'median'. Features que resultaron 100% "
        "nulas en el panel (ej. 'TAMANO_PROMEDIO_DE_GRUPO' del MEN) se descartan antes de imputar."
    )

    doc.add_heading("9.4 Selección final de variables (criterio de negocio)", level=2)
    doc.add_paragraph("Variables retenidas para el modelado, agrupadas por dimensión:")
    add_table(doc,
              ["Dimensión", "Features"],
              [
                  ("Conectividad",
                   "pct_internet_icfes, accesos_per_capita_5_16, avg_velocidad_bajada, n_proveedores"),
                  ("Pobreza",
                   "idx_privacion, pct_grupo_A, pct_rural_sisben"),
                  ("Educación (oferta)",
                   "COBERTURA_NETA, DESERCION, APROBACION"),
                  ("Tamaño y población",
                   "n_estudiantes, POBLACION_5_16"),
              ])

    # ===== 10. MODELADO RF =====
    doc.add_heading("10. Modelado supervisado: Random Forest", level=1)
    doc.add_paragraph(
        "Se entrenó un 'RandomForestRegressor' con pipeline Imputer→Scaler→RF. Se probaron 7 "
        "combinaciones de hiperparámetros (numTrees ∈ {50,100,150}, maxDepth ∈ {5,10,15}). "
        "El cuadro siguiente compara las métricas de cada configuración sobre el test set 20%."
    )
    rf_grid = find_lines("08_ml_supervisado_rf.ipynb", r"^\s+numTrees=\s*\d+\s+maxDepth=")
    if rf_grid:
        rows = []
        for ln in rf_grid:
            m = re.search(r"numTrees=\s*(\d+)\s+maxDepth=\s*(\d+).*test_R2=([\d\.]+)\s+RMSE=([\d\.]+)\s+\(([\d\.]+)s\)", ln)
            if m:
                rows.append((m.group(1), m.group(2), m.group(3), m.group(4), m.group(5)+" s"))
        if rows:
            add_table(doc, ["numTrees","maxDepth","test_R²","test_RMSE","tiempo"], rows)

    best = find_lines("08_ml_supervisado_rf.ipynb",
                       r"^Mejor modelo:", r"^\s+test_R2=")
    if best:
        doc.add_paragraph("Mejor modelo seleccionado:")
        for ln in best:
            p = doc.add_paragraph(ln); p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(10); r.font.bold = True

    add_image(doc, EVI / "rf_feature_importance.png", width_in=5.6,
              caption="Importancia de variables — Random Forest (modelo final).")

    doc.add_paragraph(
        "El modelo persistido vive en 'hdfs:///usr/proyecto/models/rf_punt_global' y puede "
        "recargarse con 'PipelineModel.load(...)' para inferencia."
    )

    # ===== 11. MODELADO KMEANS =====
    doc.add_heading("11. Modelado no supervisado: K-Means", level=1)
    doc.add_paragraph(
        "Sobre el snapshot del año más reciente (1,098 municipios), se probó k ∈ [2, 8] "
        "con métrica WCSS y coeficiente de silueta. Se eligió el k que maximiza la silueta."
    )
    km_lines = find_lines("09_ml_no_supervisado_kmeans.ipynb",
                           r"^\s+k=\d\s+WCSS=", r"^K elegido")
    rows = []
    for ln in km_lines:
        m = re.match(r"\s+k=(\d)\s+WCSS=\s*([\d\.]+)\s+silueta=([\d\.]+)", ln)
        if m: rows.append((m.group(1), m.group(2), m.group(3)))
    if rows:
        add_table(doc, ["k","WCSS","Silueta"], rows)
    add_image(doc, EVI / "kmeans_eleccion_k.png", width_in=5.5,
              caption="Selección de k: método del codo y coeficiente de silueta.")
    add_image(doc, EVI / "kmeans_clusters_scatter.png", width_in=5.6,
              caption="Tipologías municipales por K-Means (snapshot último año).")

    # ===== 12. BONO MLP =====
    doc.add_heading("12. Bono — Modelo de aprendizaje profundo (MLP)", level=1)
    doc.add_paragraph(
        "Como bono se construyó una red neuronal MLP con 'MultilayerPerceptronClassifier' de "
        "Spark MLlib. Se reformuló el problema como clasificación multi-clase (BAJO/MEDIO/ALTO) "
        "para hacerlo idóneo para la arquitectura softmax. La red corre distribuida sobre el "
        "cluster usando los mismos features que el Random Forest."
    )
    mlp_lines = find_lines("10_mlp_bono_deep_learning.ipynb",
                            r"^\s+\d+→.*acc=",
                            r"^TRAIN\s+acc=", r"^TEST\s+acc=",
                            r"Arquitectura MLP:")
    if mlp_lines:
        for ln in mlp_lines[:10]:
            p = doc.add_paragraph(ln); p.paragraph_format.left_indent = Cm(0.5)
            for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(10)
    add_image(doc, EVI / "mlp_confusion_matrix.png", width_in=4.5,
              caption="Matriz de confusión — MLP test set.")

    # ===== 13. CONCLUSIONES =====
    doc.add_heading("13. Conclusiones y recomendaciones de negocio", level=1)
    doc.add_paragraph(
        "El análisis confirma que la brecha digital es un predictor robusto del rendimiento Saber 11 "
        "municipal en Colombia. Los hallazgos más relevantes para el Ministerio:"
    )
    bullets = [
        "Existe correlación positiva clara entre acceso a internet en el hogar (declarado en ICFES) "
        "y puntaje promedio municipal. El efecto se sostiene aún controlando por privación SISBEN.",

        "La brecha rural / urbano se solapa con la brecha digital — política focalizada en municipios "
        "rurales con alta privación tendría doble dividendo (conectividad + rendimiento).",

        "El Random Forest predice el puntaje municipal con R² ≈ 0.58 usando solo 8-10 features de "
        "contexto, lo que valida que el modelo capta señal real (no ruido).",

        "K-Means identifica tipologías municipales que permiten priorizar inversión: municipios "
        "del cluster 'bajo rendimiento + alta privación' son el target natural del plan de acción.",

        "La velocidad de internet (no solo el acceso) también pesa — recomendar políticas que "
        "incentiven banda ancha de calidad, no solo conexión básica.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")

    # ===== 14. REPRODUCCION =====
    doc.add_heading("14. Cómo reproducir", level=1)
    doc.add_paragraph(
        "Repositorio del proyecto: 'proyecto/' (Mac local) y '/home/estudiante/proyecto_datos/' "
        "(cluster). Ejecución completa de la cadena:"
    )
    p = doc.add_paragraph("ssh mpde12\nbash ~/proyecto_datos/scripts/run_all.sh")
    for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(10)

    doc.add_paragraph("Para regenerar un notebook desde su spec editable:")
    p = doc.add_paragraph("python3 ~/proyecto_datos/specs/spec_08_ml_rf.py")
    for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(10)

    # ===== 15. ANEXOS =====
    doc.add_heading("15. Anexos", level=1)

    doc.add_heading("15.1 Estructura del repositorio", level=2)
    doc.add_paragraph(
        "proyecto/\n"
        "├── README.md\n"
        "├── RESULTADOS.md\n"
        "├── Reporte_Entrega2.docx          ← este documento\n"
        "├── scripts/\n"
        "│   ├── build_notebook.py         (helper nbformat con membretado)\n"
        "│   ├── common_spark.py           (SparkSession + path constants)\n"
        "│   ├── run_all.sh                (cadena de ejecución 00 → 10)\n"
        "│   ├── capture_evidence.sh       (snapshots HDFS+Spark)\n"
        "│   └── build_report.py           (genera este .docx)\n"
        "├── specs/                        (1 .py por notebook)\n"
        "├── notebooks_entrega2/           (11 .ipynb ejecutados)\n"
        "└── evidencia/                    (.png, .json, .html, .txt)"
    )
    for p in doc.paragraphs[-1].runs:
        p.font.name = "Consolas"; p.font.size = Pt(9)

    doc.add_heading("15.2 Notebooks producidos", level=2)
    add_table(doc,
              ["#", "Notebook", "Propósito"],
              [
                  ("00","00_setup_verificacion","Conexión Spark↔HDFS, sanity check"),
                  ("01","01_bronze_csv_a_parquet","Ingesta de 4 CSV → Parquet snappy"),
                  ("02","02_silver_icfes","Limpieza ICFES (filtros, binarización, geo)"),
                  ("03","03_silver_internet","Parse velocidades, normaliza códigos"),
                  ("04","04_silver_sisben","Agregación persona → municipio (avg I1..I15)"),
                  ("05","05_silver_men","Slugify cols, parse % → double"),
                  ("06","06_gold_panel_municipal","Joins (mpio × año) → mesa analítica"),
                  ("07","07_eda_y_preguntas_negocio","8 preguntas + 8 gráficos"),
                  ("08","08_ml_supervisado_rf","RF con correlación + grid de hiperparámetros"),
                  ("09","09_ml_no_supervisado_kmeans","K-Means con codo + silueta"),
                  ("10","10_mlp_bono_deep_learning","Bono: red neuronal MLP clasificador"),
              ])

    doc.add_heading("15.3 Evidencia capturada", level=2)
    doc.add_paragraph(
        "Carpeta 'proyecto/evidencia/' contiene: snapshots HTML del Spark Master UI; "
        "JSON con estado del cluster y workers; reporte 'dfsadmin' de HDFS; listado completo "
        "del layout HDFS; y los 10+ gráficos PNG generados por los notebooks de EDA, "
        "Random Forest, K-Means y MLP."
    )

    # ===== 16. BIBLIO =====
    doc.add_heading("16. Bibliografía y fuentes de datos", level=1)
    biblio = [
        "Apache Spark 3.5.2 Documentation — https://spark.apache.org/docs/3.5.2/",
        "Apache Hadoop 3.3.6 Documentation — https://hadoop.apache.org/docs/r3.3.6/",
        "Spark MLlib Programming Guide — https://spark.apache.org/docs/3.5.2/ml-guide.html",
        "ICFES Saber 11 — Resultados únicos: datos.gov.co (dataset oficial ICFES)",
        "MinTIC Colombia — Internet Fijo accesos por tecnología y segmento (datos.gov.co)",
        "DNP — SISBEN Personas (datos.gov.co)",
        "MEN — Estadísticas en Educación Preescolar, Básica y Media por Municipio (datos.gov.co)",
        "CRISP-DM — Cross-Industry Standard Process for Data Mining (Chapman et al., 2000).",
    ]
    for b in biblio:
        doc.add_paragraph(b, style="List Bullet")

    # Guardar
    doc.save(OUT)
    print(f"\nDocumento generado: {OUT}")
    print(f"Tamaño: {OUT.stat().st_size/1024:.1f} KB")


if __name__ == "__main__":
    build_report()
